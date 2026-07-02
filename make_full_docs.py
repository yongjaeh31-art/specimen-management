# make_full_docs.py — 검체관리프로그램 종합 문서 생성 (PDF + DOCX)
# 실행: python make_full_docs.py
# 출력: D:\검체관리프로그램_문서\

import os
from datetime import date
from pathlib import Path

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import (
    BaseDocTemplate, Frame, PageTemplate,
    Paragraph, Spacer, Table, TableStyle,
    HRFlowable, KeepTogether, PageBreak,
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 출력 경로 ────────────────────────────────────────────────────────────
OUT_DIR = Path("D:/검체관리프로그램_문서")
OUT_DIR.mkdir(parents=True, exist_ok=True)
PDF_PATH  = OUT_DIR / "검체관리프로그램_상세스펙_및_사용설명서.pdf"
DOCX_PATH = OUT_DIR / "검체관리프로그램_상세스펙_및_사용설명서.docx"
TODAY = date.today().strftime("%Y-%m-%d")

# ════════════════════════════════════════════════════════════════════════
#  PDF 생성
# ════════════════════════════════════════════════════════════════════════
FONT_DIR = Path("C:/Windows/Fonts")
pdfmetrics.registerFont(TTFont("Malgun",   str(FONT_DIR / "malgun.ttf")))
pdfmetrics.registerFont(TTFont("MalgunBd", str(FONT_DIR / "malgunbd.ttf")))

C_BLUE   = colors.HexColor("#1e40af")
C_LBLUE  = colors.HexColor("#dbeafe")
C_DBLUE  = colors.HexColor("#1e3a8a")
C_GRAY   = colors.HexColor("#6b7280")
C_LGRAY  = colors.HexColor("#f3f4f6")
C_GREEN  = colors.HexColor("#166534")
C_LGREEN = colors.HexColor("#dcfce7")
C_RED    = colors.HexColor("#991b1b")
C_LRED   = colors.HexColor("#fee2e2")
C_DARK   = colors.HexColor("#111827")
C_LINE   = colors.HexColor("#e5e7eb")
C_AMBER  = colors.HexColor("#92400e")
C_LAMBER = colors.HexColor("#fef3c7")
C_TEAL   = colors.HexColor("#0f766e")
C_LTEAL  = colors.HexColor("#ccfbf1")

W, H = A4


def make_styles():
    base = dict(fontName="Malgun", leading=18)
    return {
        "cover_title": ParagraphStyle("cover_title", fontName="MalgunBd", fontSize=24,
                                       textColor=colors.white, leading=32, spaceAfter=8),
        "cover_sub":   ParagraphStyle("cover_sub",   fontName="Malgun",   fontSize=12,
                                       textColor=colors.HexColor("#bfdbfe"), leading=18),
        "cover_date":  ParagraphStyle("cover_date",  fontName="Malgun",   fontSize=9,
                                       textColor=colors.HexColor("#93c5fd")),
        "title":  ParagraphStyle("title",  fontName="MalgunBd", fontSize=20, textColor=C_BLUE,
                                  spaceAfter=8),
        "h1":     ParagraphStyle("h1",     fontName="MalgunBd", fontSize=14, textColor=C_DBLUE,
                                  spaceBefore=16, spaceAfter=6),
        "h2":     ParagraphStyle("h2",     fontName="MalgunBd", fontSize=11, textColor=C_DARK,
                                  spaceBefore=10, spaceAfter=4),
        "h3":     ParagraphStyle("h3",     fontName="MalgunBd", fontSize=10, textColor=C_TEAL,
                                  spaceBefore=6,  spaceAfter=3),
        "body":   ParagraphStyle("body",   **base, fontSize=10, textColor=C_DARK, spaceAfter=4),
        "bullet": ParagraphStyle("bullet", **base, fontSize=10, textColor=C_DARK,
                                  leftIndent=16, spaceAfter=3, bulletIndent=6),
        "bullet2":ParagraphStyle("bullet2",**base, fontSize=9,  textColor=C_GRAY,
                                  leftIndent=32, spaceAfter=2),
        "note":   ParagraphStyle("note",   **base, fontSize=9,  textColor=C_AMBER),
        "code":   ParagraphStyle("code",   fontName="Malgun",   fontSize=8.5,
                                  textColor=C_DBLUE, leftIndent=12, spaceAfter=2, leading=14),
        "small":  ParagraphStyle("small",  **base, fontSize=8,  textColor=C_GRAY),
        "toc":    ParagraphStyle("toc",    **base, fontSize=10, textColor=C_DARK, spaceAfter=3),
        "toc2":   ParagraphStyle("toc2",   **base, fontSize=9,  textColor=C_GRAY,
                                  leftIndent=16, spaceAfter=2),
    }


S = make_styles()


def hr(c=C_LINE, t=0.5): return HRFlowable(width="100%", thickness=t, color=c, spaceAfter=8)
def sp(h=6): return Spacer(1, h)
def B(text): return f"<b>{text}</b>"
def BLUE(text): return f'<font color="#1e40af"><b>{text}</b></font>'


def cover_page():
    data = [
        [Paragraph("검체관리프로그램", S["cover_title"])],
        [Paragraph("상세 스펙 · 구현 방향 · 사용설명서", S["cover_sub"])],
        [sp(4)],
        [Paragraph(f"발행일: {TODAY}  |  버전: 최신 빌드", S["cover_date"])],
        [Paragraph("임상검사실 검체 소분류·배양 관리 통합 시스템", S["cover_date"])],
    ]
    t = Table(data, colWidths=[155 * mm])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), C_DBLUE),
        ("ROUNDEDCORNERS", [6]),
        ("TOPPADDING", (0, 0), (-1, 0), 22),
        ("BOTTOMPADDING", (0, -1), (-1, -1), 18),
        ("LEFTPADDING", (0, 0), (-1, -1), 18),
    ]))
    return [sp(30), t]


def section_box(text, bg=C_LBLUE, tc=C_DBLUE):
    t = Table([[Paragraph(text, ParagraphStyle("sb", fontName="MalgunBd", fontSize=13,
                                                textColor=tc, leading=20))]],
              colWidths=[165 * mm])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), bg),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ("LEFTPADDING", (0, 0), (-1, -1), 12),
    ]))
    return t


def info_table(rows, col_w=None):
    col_w = col_w or [45 * mm, 120 * mm]
    t = Table(
        [[Paragraph(B(r[0]), S["body"]), Paragraph(r[1], S["body"])] for r in rows],
        colWidths=col_w,
    )
    t.setStyle(TableStyle([
        ("BACKGROUND",   (0, 0), (0, -1), C_LGRAY),
        ("GRID",         (0, 0), (-1, -1), 0.4, C_LINE),
        ("TOPPADDING",   (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING",(0, 0), (-1, -1), 5),
        ("LEFTPADDING",  (0, 0), (-1, -1), 8),
        ("VALIGN",       (0, 0), (-1, -1), "TOP"),
    ]))
    return t


def feature_table(headers, rows, col_w=None):
    hdr = [Paragraph(B(h), ParagraphStyle("th", fontName="MalgunBd", fontSize=9,
                                            textColor=colors.white, leading=14)) for h in headers]
    body = [[Paragraph(str(c), S["small"]) for c in row] for row in rows]
    t = Table([hdr] + body, colWidths=col_w)
    t.setStyle(TableStyle([
        ("BACKGROUND",   (0, 0), (-1, 0), C_DBLUE),
        ("BACKGROUND",   (0, 1), (-1, -1), colors.white),
        ("ROWBACKGROUNDS",(0, 1), (-1, -1), [colors.white, C_LGRAY]),
        ("GRID",         (0, 0), (-1, -1), 0.4, C_LINE),
        ("TOPPADDING",   (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING",(0, 0), (-1, -1), 4),
        ("LEFTPADDING",  (0, 0), (-1, -1), 6),
        ("VALIGN",       (0, 0), (-1, -1), "TOP"),
    ]))
    return t


def add_page_num(canvas, doc):
    canvas.saveState()
    canvas.setFont("Malgun", 8)
    canvas.setFillColor(C_GRAY)
    canvas.drawCentredString(W / 2, 18 * mm, f"검체관리프로그램  —  {doc.page}")
    canvas.restoreState()


def build_pdf():
    doc = BaseDocTemplate(
        str(PDF_PATH), pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm,
        topMargin=20 * mm, bottomMargin=25 * mm,
    )
    frame = Frame(doc.leftMargin, doc.bottomMargin,
                  doc.width, doc.height, id="main")
    doc.addPageTemplates([PageTemplate(id="main", frames=[frame], onPage=add_page_num)])

    story = []

    # ── 표지 ─────────────────────────────────────────────────────────────
    story += cover_page()
    story.append(PageBreak())

    # ── 목차 ─────────────────────────────────────────────────────────────
    story.append(Paragraph("목  차", S["title"]))
    story.append(hr(C_BLUE, 1))
    toc_items = [
        ("1", "프로그램 개요 및 목적"),
        ("2", "기술 스택 및 시스템 아키텍처"),
        ("3", "데이터베이스 구조"),
        ("4", "화면별 상세 스펙"),
        ("  4.1", "접수리스트 업로드 (/upload)"),
        ("  4.2", "검체 도착·조회 (/scan)"),
        ("  4.3", "미생물 소분류 (/micro)"),
        ("  4.4", "학부별 소분류 (/micro → 비미생물)"),
        ("  4.5", "Urine 출력 (/urine)"),
        ("  4.6", "검체 찾기 (/find)"),
        ("  4.7", "누락검체 (/missing)"),
        ("  4.8", "미접수검체 (/unregistered)"),
        ("  4.9", "지사별 검체 확인 (/branch-rack)"),
        ("  4.10", "대시보드 (/dashboard)"),
        ("  4.11", "사용자 승인 (/admin/users)"),
        ("5", "구현 방향 및 설계 원칙"),
        ("6", "API 엔드포인트 목록"),
        ("7", "화면별 사용설명서"),
        ("8", "배포 및 설치 안내"),
        ("9", "GitHub 관리 안내"),
    ]
    for num, title in toc_items:
        style = S["toc"] if not num.startswith("  ") else S["toc2"]
        story.append(Paragraph(f"{num}.  {title}", style))
    story.append(PageBreak())

    # ════════════════════════════════════════════════════════════════════
    # 1. 개요 및 목적
    # ════════════════════════════════════════════════════════════════════
    story.append(section_box("1. 프로그램 개요 및 목적"))
    story.append(sp(8))
    story.append(Paragraph("1.1 개요", S["h1"]))
    story.append(Paragraph(
        "검체관리프로그램은 임상검사실에서 수신되는 검체(혈액·소변·객담·대변 등)를 "
        "접수부터 분류·배양 완료까지 일관되게 추적·관리하는 통합 웹 시스템입니다. "
        "LIS(검사정보시스템)에서 내보낸 접수 목록 Excel 파일을 업로드하면, "
        "검체를 자동으로 학부별·배양 유형별로 분류하고 랙 위치를 부여합니다. "
        "바코드 스캐너로 스캔하면 음성(TTS)으로 위치를 안내하여 분류 오류를 최소화합니다.",
        S["body"],
    ))
    story.append(sp(4))

    story.append(Paragraph("1.2 개발 배경 및 목적", S["h1"]))
    story.append(info_table([
        ("도입 배경", "수작업 검체 분류로 인한 오류(랙 위치 혼동, 누락) 및 업무 지연 개선"),
        ("핵심 목적", "접수번호 바코드 스캔 → 자동 위치 안내 → 작업 표준화"),
        ("적용 환경", "단일 PC 또는 네트워크 공유(서버 1대 + 클라이언트 N대) 환경"),
        ("주 사용자", "임상검사실 전담 직원 (미생물팀, 혈액·화학·지사 검체 담당)"),
        ("운영 시간", "24시간 운영 (야간 19:00 ~ 익일 09:00 별도 세션 지원)"),
    ]))
    story.append(sp(6))

    story.append(Paragraph("1.3 주요 기능 목록", S["h1"]))
    feat_rows = [
        ("LIS 파일 업로드", "Excel(xls/xlsx) 접수 목록 → DB 일괄 등록, 중복 방지"),
        ("검체 도착 스캔", "바코드 스캔 → 도착 시각 기록, 학부·분류코드 표시"),
        ("미생물 소분류", "배양 유형별 랙·칸 번호 자동 부여, TTS 음성 안내"),
        ("학부별 소분류", "비미생물(화학·혈액·요경검 등) 랙 위치 자동 배정"),
        ("지사별 검체 확인", "분류코드 선택 → 워크리스트 생성 → 바코드 스캔 확인"),
        ("Urine 출력", "접수 목록에서 Urine 해당 행 마킹 Excel 다운로드"),
        ("검체 찾기", "접수번호 검색 → 랙·칸 번호, 도착 여부, 검사명 일괄 표시"),
        ("누락검체 조회", "도착하지 않은 검체 목록 필터링·출력"),
        ("미접수검체 조회", "LIS에 없는 검체 스캔 기록 목록 표시"),
        ("대시보드", "실시간 도착률, 학부별 진행 현황, 오늘 업로드 현황"),
        ("사용자 관리", "회원가입 승인, 계정 활성/비활성 관리"),
        ("TTS 음성 안내", "Web Speech API, 음성·속도 선택, 기기 간 설정 공유"),
        ("네트워크 공유", "bat 파일 1개로 서버 공유, 다른 PC에서 IP 접속"),
        ("출근 기준 리셋", "출근 시각(기본 19:00) 기준 소분류 자동 초기화"),
    ]
    story.append(feature_table(
        ["기능", "설명"],
        feat_rows,
        col_w=[50 * mm, 115 * mm],
    ))
    story.append(PageBreak())

    # ════════════════════════════════════════════════════════════════════
    # 2. 기술 스택 및 아키텍처
    # ════════════════════════════════════════════════════════════════════
    story.append(section_box("2. 기술 스택 및 시스템 아키텍처"))
    story.append(sp(8))

    story.append(Paragraph("2.1 기술 스택", S["h1"]))
    story.append(feature_table(
        ["구분", "기술", "버전/비고"],
        [
            ("Backend",  "FastAPI",                 "0.115 — 비동기 REST API 서버"),
            ("ORM",      "SQLAlchemy",              "2.0 — Mapped/mapped_column 선언형"),
            ("DB",       "SQLite (WAL 모드)",        "파일 기반, 다중 읽기·단일 쓰기 동시성"),
            ("Frontend", "Jinja2 템플릿",            "HTML 서버 사이드 렌더링"),
            ("UI 프레임", "Bootstrap 5.3",           "반응형 레이아웃, 모달·콜랩스 컴포넌트"),
            ("음성",     "Web Speech API",           "브라우저 내장 TTS, SpeechSynthesis"),
            ("Excel",    "openpyxl 3.1 / xlrd",     "xlsx 읽기·쓰기, xls 읽기"),
            ("PDF",      "reportlab",               "문서 자동 생성"),
            ("Word",     "python-docx 1.2",         "사용설명서 DOCX 생성"),
            ("런타임",   "Python 3.11 embedded",     "별도 설치 없이 동봉 (약 65 MB)"),
            ("실행파일", "검체관리프로그램.exe",        "PyInstaller 빌드, 서버 자동 시작"),
            ("CI/CD",    "GitHub Actions",           "태그 푸시 → 릴리즈 ZIP 2종 자동 생성"),
        ],
        col_w=[28 * mm, 45 * mm, 92 * mm],
    ))
    story.append(sp(8))

    story.append(Paragraph("2.2 시스템 아키텍처", S["h1"]))
    story.append(Paragraph(
        "단일 서버 구조: FastAPI 서버(uvicorn)가 SQLite 파일을 직접 읽고 쓰며, "
        "여러 클라이언트 브라우저에서 동시 접속합니다. "
        "네트워크 공유 모드에서는 서버 PC의 IP:8000으로 다른 PC가 접속합니다.",
        S["body"],
    ))
    arch_rows = [
        ["클라이언트 브라우저", "↕ HTTP/REST", "FastAPI 서버 (uvicorn)", "↕ SQLAlchemy ORM", "SQLite DB"],
    ]
    t = Table(arch_rows, colWidths=[38 * mm, 22 * mm, 48 * mm, 30 * mm, 27 * mm])
    t.setStyle(TableStyle([
        ("BACKGROUND",   (0, 0), (0, -1), C_LGREEN),
        ("BACKGROUND",   (2, 0), (2, -1), C_LBLUE),
        ("BACKGROUND",   (4, 0), (4, -1), C_LGRAY),
        ("GRID",         (0, 0), (-1, -1), 0.4, C_LINE),
        ("FONTNAME",     (0, 0), (-1, -1), "Malgun"),
        ("FONTSIZE",     (0, 0), (-1, -1), 9),
        ("ALIGN",        (0, 0), (-1, -1), "CENTER"),
        ("VALIGN",       (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING",   (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING",(0, 0), (-1, -1), 8),
    ]))
    story.append(t)
    story.append(sp(8))

    story.append(Paragraph("2.3 파일 구조", S["h1"]))
    file_rows = [
        ("app/main.py",              "FastAPI 앱 진입점, 라우터 등록, DB 초기화, 스키마 마이그레이션"),
        ("app/models.py",            "SQLAlchemy 테이블 모델 (13개 테이블)"),
        ("app/database.py",          "DB 엔진 생성, WAL 모드 설정, get_db 의존성"),
        ("app/routers/*.py",         "REST API 라우터 (imports, scans, micro, reports, auth, branch_rack)"),
        ("app/services/*.py",        "비즈니스 로직 서비스 레이어"),
        ("app/templates/*.html",     "Jinja2 HTML 템플릿 (11개 화면)"),
        ("app/static/css/styles.css","공통 CSS (네비게이션, 카드, 배지 등)"),
        ("app/static/js/app.js",     "공통 JS (실시간 시계, 로그인 상태 관리)"),
        ("launcher.py",              "uvicorn 서버 시작 + 브라우저 자동 오픈"),
        ("검체관리프로그램.exe",       "PyInstaller 단일 실행파일 (launcher.py 래퍼)"),
        ("시작.bat",                  "서버 시작 배치 파일 (python 경로 지정)"),
        ("네트워크공유.bat",           "0.0.0.0 바인딩으로 네트워크 공유 시작"),
        ("배포/",                     "배포 ZIP 아카이브 (날짜별 보관)"),
    ]
    story.append(feature_table(
        ["파일/경로", "설명"],
        file_rows,
        col_w=[55 * mm, 110 * mm],
    ))
    story.append(PageBreak())

    # ════════════════════════════════════════════════════════════════════
    # 3. 데이터베이스 구조
    # ════════════════════════════════════════════════════════════════════
    story.append(section_box("3. 데이터베이스 구조"))
    story.append(sp(8))
    story.append(Paragraph(
        "SQLite 단일 파일 DB (specimen_routing.db). WAL 모드로 다중 동시 읽기를 지원합니다. "
        "스키마 변경 시 app/main.py의 ensure_schema_columns()에서 ALTER TABLE로 자동 추가합니다.",
        S["body"],
    ))
    story.append(sp(6))

    db_tables = [
        ("import_batches",                    "업로드 배치 이력 (파일명, 차수, 최종 여부, 행 수)"),
        ("orders",                            "검체 접수 목록 (접수번호, 환자명, 나이, 검체명, 병원명, 접수일)"),
        ("order_tests",                       "검사 항목 (접수번호당 N건, 검사명, 학부, 검체명 override)"),
        ("specimen_arrivals",                 "검체 도착 기록 (도착 시각, 도착자, PC명, 미접수 여부)"),
        ("scan_logs",                         "전체 스캔 이력 로그 (스캔 유형, 결과, 오퍼레이터)"),
        ("department_routes",                 "학부별 라우팅 결과 캐시 (접수번호 → 학부)"),
        ("routing_rules",                     "검사명 → 학부 매핑 규칙 (시작 시 시드 데이터)"),
        ("culture_rules",                     "미생물 배양 유형별 prefix·순번 관리"),
        ("micro_culture_plans",               "업로드 시 사전 생성되는 미생물 배양 예정 목록"),
        ("micro_culture_assignments",         "실제 스캔 완료된 배양 배정 (랙·칸·코드)"),
        ("department_subcategory_rules",      "학부별 소분류(비미생물) 유형·prefix·순번"),
        ("department_subcategory_assignments","학부별 소분류 배정 결과 (랙·칸·위치코드)"),
        ("branch_rack_sessions",              "지사별 검체 확인 세션 (PC별, 분류코드 JSON, 총 건수)"),
        ("branch_rack_items",                 "세션 내 개별 검체 (랙번호, 칸번호, 스캔 여부, 시각)"),
        ("app_users",                         "사용자 계정 (아이디, 표시이름, 해시 비밀번호, 승인 여부)"),
        ("registered_pcs",                    "등록된 PC 목록"),
        ("app_settings",                      "전역 설정 키-값 (예: 출근 기준 시각, 평일/토요일 구분)"),
    ]
    story.append(feature_table(
        ["테이블명", "역할"],
        db_tables,
        col_w=[65 * mm, 100 * mm],
    ))
    story.append(PageBreak())

    # ════════════════════════════════════════════════════════════════════
    # 4. 화면별 상세 스펙
    # ════════════════════════════════════════════════════════════════════
    story.append(section_box("4. 화면별 상세 스펙"))
    story.append(sp(8))

    # 4.1 업로드
    story.append(Paragraph("4.1  접수리스트 업로드  (/upload)", S["h1"]))
    story.append(info_table([
        ("경로",     "/upload"),
        ("파일",     "app/templates/upload.html  /  app/routers/imports.py"),
        ("목적",     "LIS에서 내보낸 Excel 파일(xls/xlsx)을 DB에 일괄 등록"),
        ("주요 입력", "• Excel 파일 첨부\n• 차수 선택 (1차 / 2차 / 3차)\n• 최종 업로드 체크박스\n• 출근 기준 (평일 / 토요일)"),
        ("주요 출력", "• 업로드 성공/실패 토스트\n• 등록 행 수, 배치 이력 테이블\n• 현재 등록된 접수 총 건수"),
        ("특이사항", "• 동일 접수번호 중복 업로드 시 업데이트(merge)\n"
                    "• 1차·2차·3차 누적 업로드 지원 (신규만 추가)\n"
                    "• 최종 업로드 표시 → 해당 배치만 is_final=True\n"
                    "• 업로드 동시에 미생물 예정 목록(MicroCulturePlan) 자동 생성\n"
                    "• '파일 삭제 및 초기화' → 모든 데이터 일괄 삭제\n"
                    "  (지사별 검체 확인 세션 포함)"),
    ]))
    story.append(sp(8))

    # 4.2 검체 도착
    story.append(Paragraph("4.2  검체 도착·조회  (/scan)", S["h1"]))
    story.append(info_table([
        ("경로",     "/scan"),
        ("파일",     "app/templates/scan.html  /  app/routers/scans.py"),
        ("목적",     "바코드 스캔으로 검체 도착 기록 및 학부 표시"),
        ("주요 입력", "• 바코드 스캐너 (접수번호 7자리 또는 15자리 전체)\n• 검체 카테고리 선택 필터"),
        ("주요 출력", "• 도착 카드: 환자명, 나이, 병원명, 검사명 목록\n• 학부별 색상 배지\n• 재스캔 경고"),
        ("특이사항", "• 15자리 → 7자리 접수번호 자동 변환 (barcode_service)\n"
                    "• 미접수검체(DB에 없는 접수번호) 별도 기록\n"
                    "• 알리쿼트·전달 필요 여부 배지 표시"),
    ]))
    story.append(sp(8))

    # 4.3 미생물
    story.append(Paragraph("4.3  미생물 소분류  (/micro → 미생물 탭)", S["h1"]))
    story.append(info_table([
        ("경로",     "/micro"),
        ("파일",     "app/templates/micro.html  /  app/routers/micro.py  /  app/services/micro_service.py"),
        ("목적",     "미생물 배양 검체를 유형별 랙에 자동 배정하고 스캔으로 확인"),
        ("배양 유형", "일반세균, 혈액배양, 결핵, 진균, 혐기, 연관(MBI/PCR 등) 등 복수 선택"),
        ("주요 입력", "• 배양 유형 버튼 선택 (복수 가능)\n• 바코드 스캔\n• 출근 기준 시각 설정"),
        ("주요 출력", "• 랙 번호 + 칸 번호 (대형 폰트)\n• TTS 음성 ('1번랙 5번칸')\n"
                    "• 랙 그리드 시각화 (10열 컬러 셀)"),
        ("특이사항", "• 업로드 시 접수번호 오름차순으로 예정 순번 사전 부여\n"
                    "• 스캔 순서와 무관하게 LIS 순번 위치 유지\n"
                    "• 출근 기준 시각 이후 첫 스캔 → 랙 자동 리셋\n"
                    "• 랙 만료 시 다음 랙 자동 시작\n"
                    "• 연관검사(MBI 등) 별도 배양 유형으로 분리 관리"),
    ]))
    story.append(sp(8))

    # 4.4 학부별
    story.append(Paragraph("4.4  학부별 소분류  (/micro → 비미생물 탭)", S["h1"]))
    story.append(info_table([
        ("경로",     "/micro (탭 전환)"),
        ("파일",     "app/templates/micro.html  /  app/services/subdivision_service.py"),
        ("목적",     "혈액·화학·요경검·분자진단 등 비미생물 검체를 소분류 기준에 따라 랙 배정"),
        ("주요 기능", "• 학부 선택 → 소분류 선택 → 바코드 스캔\n"
                    "• 랙 그리드 10열 시각화\n"
                    "• TTS 음성 안내\n"
                    "• 소분류 완료·미완료 현황 배지"),
        ("위치 코드", "prefix + sequence_no 조합, 예: BCH-001, HEM-003"),
    ]))
    story.append(sp(8))

    # 4.5 Urine
    story.append(Paragraph("4.5  Urine 출력  (/urine)", S["h1"]))
    story.append(info_table([
        ("경로",     "/urine"),
        ("파일",     "app/templates/urine.html  /  app/routers/imports.py (urine-mark 엔드포인트)"),
        ("목적",     "접수 목록 Excel에서 Urine Random + OC&S 해당 행에 마킹 후 다운로드"),
        ("처리 내용", "• 검사명 셀 앞 ▣ 기호 삽입\n• 인접 셀에 ● 기호 삽입\n• 수정된 Excel 다운로드"),
    ]))
    story.append(sp(8))

    # 4.6 검체 찾기
    story.append(Paragraph("4.6  검체 찾기  (/find)", S["h1"]))
    story.append(info_table([
        ("경로",     "/find"),
        ("파일",     "app/templates/find.html"),
        ("목적",     "접수번호로 해당 검체의 소분류 위치·도착 여부·검사명 한눈에 확인"),
        ("주요 출력", "• 미생물 배양 위치 (랙·칸·코드)\n• 학부별 소분류 위치\n• 도착 여부 및 시각\n• 전체 검사명 목록"),
    ]))
    story.append(sp(8))

    # 4.7 누락
    story.append(Paragraph("4.7  누락검체  (/missing)", S["h1"]))
    story.append(info_table([
        ("경로",     "/missing"),
        ("파일",     "app/templates/missing.html"),
        ("목적",     "도착 기록이 없는 검체(미도착)를 필터링하여 목록 표시"),
        ("필터",     "학부, 도착 여부, 검색어"),
        ("출력",     "접수번호, 환자명, 검사명, 병원명, 등록 시각"),
    ]))
    story.append(sp(8))

    # 4.8 미접수
    story.append(Paragraph("4.8  미접수검체  (/unregistered)", S["h1"]))
    story.append(info_table([
        ("경로",     "/unregistered"),
        ("파일",     "app/templates/unregistered.html"),
        ("목적",     "LIS에 등록되지 않은 접수번호가 스캔된 경우 목록 표시"),
        ("활용",     "LIS 누락 접수 확인, 오스캔 점검"),
    ]))
    story.append(sp(8))

    # 4.9 지사별
    story.append(Paragraph("4.9  지사별 검체 확인  (/branch-rack)", S["h1"]))
    story.append(info_table([
        ("경로",     "/branch-rack"),
        ("파일",     "app/templates/branch_rack.html  /  app/services/branch_rack_service.py"),
        ("목적",     "지사코드(접수번호 앞 2자리)별로 검체를 랙에 배치하고 바코드 스캔으로 확인"),
        ("주요 흐름", "① PC명 입력 → ② 분류코드 선택 → ③ 준비 시작 → ④ 바코드 스캔"),
        ("준비 시작", "• 동일 코드 재호출 → 신규 검체만 추가 (기존 위치·스캔상태 유지)\n"
                    "• 코드 변경 → 세션 초기화 후 재생성"),
        ("랙 계산",  "sort_no 기준: rack_no = (sort_no-1)÷50+1, rack_position = (sort_no-1)%50+1"),
        ("TTS",      "스캔 시 '1번랙 5번' 음성 (접수번호 생략, 랙·위치만 안내)"),
        ("미확인 목록","'미확인 N건' 버튼 클릭 → 접이식 패널 표시, 인쇄 가능\n"
                    "  열: 순번, 접수번호, 랙, 자리, 환자명, 검체명, 병원명"),
        ("랙 그리드", "10열 색상 셀 (회색=빈자리, 녹색=스캔완료, 주황=방금 스캔)\n"
                    "  3.5초 후 주황 → 녹색으로 자동 전환"),
        ("초기화",   "'파일 삭제 및 초기화' 버튼 → 세션 및 아이템 모두 삭제"),
        ("폴링",     "3초 주기 서버 상태 갱신 (다중 PC 동시 스캔 지원)"),
    ]))
    story.append(sp(8))

    # 4.10 대시보드
    story.append(Paragraph("4.10  대시보드  (/dashboard)", S["h1"]))
    story.append(info_table([
        ("경로",     "/dashboard"),
        ("파일",     "app/templates/dashboard.html"),
        ("목적",     "금일 업무 현황 요약 (업로드·도착·소분류 진행률)"),
        ("표시 항목", "• 총 접수 건수, 도착 건수, 도착률\n• 미생물 배양 진행 현황\n• 학부별 소분류 진행 현황\n• 업로드 배치 이력"),
    ]))
    story.append(sp(8))

    # 4.11 사용자 승인
    story.append(Paragraph("4.11  사용자 승인  (/admin/users)", S["h1"]))
    story.append(info_table([
        ("경로",     "/admin/users"),
        ("파일",     "app/templates/admin_users.html  /  app/routers/auth.py"),
        ("목적",     "회원가입 신청 계정 승인, 활성/비활성 관리"),
        ("기본 계정", "admin1234 / admin1234 (자동 생성)"),
        ("인증 방식", "localStorage에 JWT 유사 세션 저장, 로그인 만료 없음"),
    ]))
    story.append(PageBreak())

    # ════════════════════════════════════════════════════════════════════
    # 5. 구현 방향 및 설계 원칙
    # ════════════════════════════════════════════════════════════════════
    story.append(section_box("5. 구현 방향 및 설계 원칙"))
    story.append(sp(8))

    principles = [
        ("무설치 실행", "Python 3.11 임베디드 런타임과 모든 패키지를 배포 ZIP에 동봉. "
                       "사용자는 파이썬 설치 없이 시작.bat 또는 검체관리프로그램.exe 더블클릭만으로 실행."),
        ("단순 배포",   "GitHub Actions가 main 브랜치 태그 푸시 시 설치용 ZIP(전체)과 "
                       "업데이트용 ZIP(app/ 폴더만)을 자동 생성하여 Releases에 업로드."),
        ("데이터 영속성","SQLite WAL 모드: 다중 브라우저 동시 읽기 허용, 단일 쓰기 락. "
                        "DB 파일이 곧 백업 대상 — 1개 파일 복사로 전체 데이터 보존."),
        ("누적 업로드", "1차→2차→3차 접수 목록 순차 업로드 시 신규 접수번호만 추가 "
                       "(ON CONFLICT IGNORE 방식). 기존 스캔 기록 보존."),
        ("랙 위치 사전 부여", "업로드 시 즉시 접수번호 오름차순으로 MicroCulturePlan에 순번 확정. "
                             "스캔 순서와 무관하게 동일 위치 재현 가능."),
        ("폴링 기반 실시간성","3초 주기 REST API 폴링으로 다중 PC 간 상태 동기화. "
                            "WebSocket 의존성 제거, 단순 HTTP 서버로 유지."),
        ("TTS 표준화",  "Web Speech API 사용. 음성·속도 설정은 localStorage에 저장하여 "
                       "화면 이동 시에도 유지. 미생물·지사별 화면 간 설정 공유."),
        ("서비스 레이어 분리","라우터(HTTP)와 서비스(비즈니스 로직)를 분리. "
                            "라우터는 의존성 주입(get_db)과 HTTP 처리만 담당."),
        ("점진적 스키마", "ensure_schema_columns()가 서버 시작 시 누락 컬럼을 ALTER TABLE로 추가. "
                        "DB 파일을 유지한 채 프로그램만 업데이트 가능."),
        ("보안",        "비밀번호 bcrypt 해싱, 관리자 승인 후 로그인 허용. "
                       "모든 화면 localStorage 세션 확인 → 미인증 시 /login 리다이렉트."),
    ]
    for i, (title, desc) in enumerate(principles, 1):
        story.append(KeepTogether([
            Paragraph(f"{i}.  {title}", S["h2"]),
            Paragraph(desc, S["body"]),
            sp(4),
        ]))
    story.append(PageBreak())

    # ════════════════════════════════════════════════════════════════════
    # 6. API 엔드포인트
    # ════════════════════════════════════════════════════════════════════
    story.append(section_box("6. API 엔드포인트 목록"))
    story.append(sp(8))
    api_rows = [
        ("POST", "/api/import/orders",              "Excel 업로드 → 접수 목록 등록"),
        ("POST", "/api/import/reset-all",           "전체 데이터 초기화"),
        ("POST", "/api/import/urine-mark",          "Urine 마킹 Excel 다운로드"),
        ("GET",  "/api/import/worklist",            "워크리스트 Excel 다운로드"),
        ("POST", "/api/scans/arrive",               "검체 도착 스캔"),
        ("GET",  "/api/scans/summary",              "도착 현황 요약"),
        ("POST", "/api/micro/scan",                 "미생물 소분류 스캔"),
        ("GET",  "/api/micro/session",              "현재 미생물 세션 조회"),
        ("POST", "/api/micro/reset",                "미생물 세션 리셋"),
        ("GET",  "/api/micro/plans",                "미생물 예정 목록 조회"),
        ("POST", "/api/micro/subdivision/scan",     "학부별 소분류 스캔"),
        ("GET",  "/api/micro/subdivision/session",  "학부별 소분류 세션 조회"),
        ("GET",  "/api/reports/missing",            "누락검체 목록"),
        ("GET",  "/api/reports/unregistered",       "미접수검체 목록"),
        ("GET",  "/api/reports/dashboard",          "대시보드 통계"),
        ("POST", "/api/branch-rack/prepare",        "지사별 검체 확인 세션 준비"),
        ("GET",  "/api/branch-rack/session",        "지사별 세션 조회"),
        ("POST", "/api/branch-rack/scan",           "지사별 바코드 스캔"),
        ("GET",  "/api/branch-rack/codes",          "사용 가능한 분류코드 목록"),
        ("POST", "/api/auth/register",              "회원가입"),
        ("POST", "/api/auth/login",                 "로그인"),
        ("GET",  "/api/auth/users",                 "사용자 목록 (관리자)"),
        ("POST", "/api/auth/approve/{id}",          "사용자 승인 (관리자)"),
    ]
    story.append(feature_table(
        ["메서드", "경로", "설명"],
        api_rows,
        col_w=[18 * mm, 72 * mm, 75 * mm],
    ))
    story.append(PageBreak())

    # ════════════════════════════════════════════════════════════════════
    # 7. 사용설명서
    # ════════════════════════════════════════════════════════════════════
    story.append(section_box("7. 화면별 사용설명서"))
    story.append(sp(8))

    story.append(Paragraph("7.1 시작 및 로그인", S["h1"]))
    steps_login = [
        ("시작",       "시작.bat (또는 검체관리프로그램.exe) 더블클릭 → 콘솔 창 열림 → 브라우저 자동 오픈"),
        ("주소",       "http://127.0.0.1:8000 (로컬) 또는 http://서버IP:8000 (네트워크)"),
        ("로그인",     "아이디 / 비밀번호 입력 → 로그인 버튼 (초기 계정: admin1234 / admin1234)"),
        ("세션 유지",  "localStorage에 저장 — 브라우저 탭 닫아도 유지, 로그아웃 버튼으로 해제"),
        ("신규 계정",  "회원가입 → 관리자(/admin/users)가 승인 후 사용 가능"),
    ]
    story.append(info_table(steps_login, col_w=[30 * mm, 135 * mm]))
    story.append(sp(8))

    story.append(Paragraph("7.2 접수리스트 업로드", S["h1"]))
    story.append(Paragraph("① 접수리스트 업로드 화면 이동", S["h3"]))
    story.append(Paragraph("② 파일 선택 → LIS에서 내보낸 Excel 파일(.xls/.xlsx) 선택", S["body"]))
    story.append(Paragraph("③ 차수 선택: 당일 처음이면 1차, 추가 접수가 있으면 2차·3차 선택", S["body"]))
    story.append(Paragraph("④ 오늘 마지막 업로드라면 '최종 업로드' 체크", S["body"]))
    story.append(Paragraph("⑤ 평일/토요일 선택 후 업로드 버튼 클릭", S["body"]))
    story.append(Paragraph(
        "⑥ 성공 시 '○○건 등록 완료' 토스트 표시 → 배치 이력 테이블 갱신", S["body"]))
    story.append(Paragraph(
        "※ 동일 접수번호 중복 업로드 시 해당 건은 스킵(중복 오류 없음)", S["note"]))
    story.append(sp(8))

    story.append(Paragraph("7.3 미생물 소분류 스캔", S["h1"]))
    story.append(Paragraph("① /micro 화면 이동 → 미생물 탭 선택", S["h3"]))
    story.append(Paragraph("② 배양 유형 버튼 선택 (복수 선택 가능 — 예: 일반세균 + 혈액배양)", S["body"]))
    story.append(Paragraph("③ TTS 토글 ON → 음성 안내 활성화 (음성·속도 선택)", S["body"]))
    story.append(Paragraph("④ 바코드 스캐너로 검체 스캔 → '2번랙 15번칸' 음성 출력 + 화면 표시", S["body"]))
    story.append(Paragraph("⑤ 랙 그리드에서 해당 위치(주황색 점멸)를 확인 후 검체 꽂기", S["body"]))
    story.append(Paragraph("⑥ 랙이 가득 차면 다음 랙 번호로 자동 이동", S["body"]))
    story.append(Paragraph(
        "※ 출근 기준 시각(기본 19:00) 이후 처음 스캔 시 랙 자동 리셋", S["note"]))
    story.append(sp(8))

    story.append(Paragraph("7.4 지사별 검체 확인", S["h1"]))
    story.append(Paragraph("① /branch-rack 화면 이동", S["h3"]))
    story.append(Paragraph("② PC명 입력 (해당 PC 식별용, 예: 지사1_PC, 야간검사실_PC)", S["body"]))
    story.append(Paragraph("③ 분류코드 버튼 선택 (접수번호 앞 2자리 — 복수 선택 가능)", S["body"]))
    story.append(Paragraph("④ '준비 시작' 버튼 클릭 → 해당 코드 검체 목록 생성 + 랙 배정", S["body"]))
    story.append(Paragraph("⑤ 바코드 스캔 → '1번랙 5번' 음성 + 랙 그리드 강조", S["body"]))
    story.append(Paragraph("⑥ '미확인 N건' 버튼 클릭 → 미스캔 검체 목록 펼침 (인쇄 가능)", S["body"]))
    story.append(Paragraph(
        "※ 워크리스트 추가(2차·3차) 후 동일 코드로 '준비 시작' 재클릭 → 신규 검체만 추가, 기존 스캔 상태 유지",
        S["note"]))
    story.append(Paragraph(
        "※ 코드를 변경하면 기존 세션 삭제 후 새로 생성",
        S["note"]))
    story.append(sp(8))

    story.append(Paragraph("7.5 검체 찾기", S["h1"]))
    story.append(Paragraph("① /find 화면 이동", S["h3"]))
    story.append(Paragraph("② 검색창에 접수번호 입력 (바코드 스캔 또는 직접 입력)", S["body"]))
    story.append(Paragraph("③ 결과 카드에서 도착 여부, 미생물 랙·칸, 학부별 위치 코드 확인", S["body"]))
    story.append(sp(8))

    story.append(Paragraph("7.6 TTS 설정", S["h1"]))
    story.append(info_table([
        ("음성 선택", "드롭다운에서 시스템 음성 선택 (Microsoft 서현 권장)"),
        ("속도 조절", "슬라이더 0.5(느림) ~ 2.0(빠름), 기본 1.0"),
        ("설정 저장", "localStorage에 저장 — 미생물·지사별 화면 공유"),
        ("음소거",    "TTS 토글 버튼으로 즉시 ON/OFF"),
    ]))
    story.append(sp(8))

    story.append(Paragraph("7.7 네트워크 공유", S["h1"]))
    story.append(info_table([
        ("서버 측",   "네트워크공유.bat 실행 → 0.0.0.0:8000 바인딩 시작"),
        ("클라이언트", "브라우저에서 http://서버IP:8000 접속"),
        ("방화벽",    "서버 PC 방화벽에서 TCP 8000 인바운드 허용 필요"),
        ("동시 접속", "복수 PC 동시 접속 가능, 폴링으로 실시간 상태 공유"),
    ]))
    story.append(PageBreak())

    # ════════════════════════════════════════════════════════════════════
    # 8. 배포 및 설치
    # ════════════════════════════════════════════════════════════════════
    story.append(section_box("8. 배포 및 설치 안내"))
    story.append(sp(8))

    story.append(Paragraph("8.1 배포 파일 종류", S["h1"]))
    story.append(feature_table(
        ["파일명", "대상", "내용"],
        [
            ("검체관리프로그램_설치_vX.X.X.zip", "처음 설치하는 PC",
             "전체 파일 포함 (python 폴더, app 폴더, bat 파일, exe, 문서)"),
            ("검체관리프로그램_업데이트_vX.X.X.zip", "이미 설치된 PC",
             "app 폴더만 포함 — 기존 app 폴더에 덮어쓰기"),
        ],
        col_w=[60 * mm, 38 * mm, 67 * mm],
    ))
    story.append(sp(6))

    story.append(Paragraph("8.2 설치 절차 (신규)", S["h1"]))
    story.append(Paragraph("① 설치 ZIP 다운로드 및 원하는 경로에 압축 해제", S["body"]))
    story.append(Paragraph("② 시작.bat 더블클릭 → 서버 시작 + 브라우저 자동 오픈", S["body"]))
    story.append(Paragraph("③ 로그인 화면에서 admin1234 / admin1234로 첫 로그인", S["body"]))
    story.append(Paragraph("④ (선택) 네트워크 공유가 필요하면 네트워크공유.bat 사용", S["body"]))
    story.append(sp(6))

    story.append(Paragraph("8.3 업데이트 절차", S["h1"]))
    story.append(Paragraph("① 업데이트 ZIP 다운로드", S["body"]))
    story.append(Paragraph("② 기존 서버 종료 (콘솔 창 닫기 또는 작업 관리자)", S["body"]))
    story.append(Paragraph("③ 업데이트 ZIP의 app 폴더를 기존 설치 경로의 app 폴더에 덮어쓰기", S["body"]))
    story.append(Paragraph("④ 시작.bat 재실행 → DB 파일(specimen_routing.db)은 유지됨", S["body"]))
    story.append(sp(6))

    story.append(Paragraph("8.4 데이터 백업", S["h1"]))
    story.append(Paragraph(
        "specimen_routing.db 파일 하나를 복사하면 모든 데이터 백업 완료. "
        "서버 실행 중에도 복사 가능(WAL 모드 덕분에 읽기 안전).",
        S["body"]))
    story.append(PageBreak())

    # ════════════════════════════════════════════════════════════════════
    # 9. GitHub 관리
    # ════════════════════════════════════════════════════════════════════
    story.append(section_box("9. GitHub 관리 안내"))
    story.append(sp(8))

    story.append(Paragraph("9.1 커밋 및 릴리즈 절차", S["h1"]))
    story.append(Paragraph("소스 코드 변경 후 GitHub에 커밋·태그 푸시하면 배포 ZIP이 자동 생성됩니다.", S["body"]))
    story.append(sp(4))
    cmd_lines = [
        "# 1. 변경 사항 확인",
        "git status",
        "",
        "# 2. 스테이징",
        "git add -A",
        "",
        "# 3. 커밋",
        'git commit -m "feat: 변경 내용 요약"',
        "",
        "# 4. 푸시",
        "git push origin main",
        "",
        "# 5. 새 버전 태그 → GitHub Actions 자동 릴리즈",
        "git tag v1.5.0",
        "git push origin v1.5.0",
    ]
    for line in cmd_lines:
        story.append(Paragraph(line if line else " ", S["code"]))
    story.append(sp(6))

    story.append(Paragraph("9.2 GitHub Actions 워크플로우", S["h1"]))
    story.append(Paragraph(
        ".github/workflows/release.yml이 태그 푸시 이벤트를 트리거합니다. "
        "워크플로우는 설치 ZIP(전체)과 업데이트 ZIP(app 폴더)을 빌드하여 "
        "해당 태그의 GitHub Releases에 자동 업로드합니다.",
        S["body"]))
    story.append(sp(6))

    story.append(Paragraph("9.3 로컬 배포 파일 수동 생성", S["h1"]))
    story.append(Paragraph(
        "GitHub 없이 로컬에서 배포 ZIP을 만들 때는 빌드_배포파일.ps1 또는 빌드_배포파일.bat을 실행합니다. "
        "결과는 배포\\ 폴더에 날짜 포함 파일명으로 저장됩니다.",
        S["body"]))

    story.append(sp(12))
    story.append(hr(C_BLUE, 1))
    story.append(Paragraph(
        f"본 문서는 {TODAY} 기준으로 생성되었습니다. "
        "프로그램 업데이트 시 내용이 달라질 수 있습니다.",
        S["small"],
    ))

    doc.build(story)
    print(f"PDF 생성 완료: {PDF_PATH}")


# ════════════════════════════════════════════════════════════════════════
#  DOCX 생성
# ════════════════════════════════════════════════════════════════════════

def set_cell_bg(cell, hex_color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "맑은 고딕"
        run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf) if level == 1 else RGBColor(0x11, 0x18, 0x27)
    return p


def add_para(doc, text, indent=0, bold=False, size=10, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "맑은 고딕"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        set_cell_bg(cell, "1e3a8a")
        for run in cell.paragraphs[0].runs:
            run.font.name = "맑은 고딕"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
            run.bold = True
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        bg = "f3f4f6" if ri % 2 == 0 else "ffffff"
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            for run in cell.paragraphs[0].runs:
                run.font.name = "맑은 고딕"
                run.font.size = Pt(9)
    if col_widths:
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                if ci < len(col_widths):
                    cell.width = Cm(col_widths[ci])
    return table


def build_docx():
    doc = Document()

    # 페이지 여백
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # 제목
    title = doc.add_heading("검체관리프로그램", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = "맑은 고딕"
        run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
        run.font.size = Pt(26)

    sub = doc.add_paragraph("상세 스펙 · 구현 방향 · 사용설명서")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.name = "맑은 고딕"
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

    date_p = doc.add_paragraph(f"발행일: {TODAY}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in date_p.runs:
        run.font.name = "맑은 고딕"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

    doc.add_page_break()

    # 1. 개요
    add_heading(doc, "1. 프로그램 개요 및 목적")
    add_para(doc, "1.1 개요", bold=True, size=11)
    add_para(doc, (
        "검체관리프로그램은 임상검사실에서 수신되는 검체(혈액·소변·객담·대변 등)를 "
        "접수부터 분류·배양 완료까지 일관되게 추적·관리하는 통합 웹 시스템입니다. "
        "LIS에서 내보낸 Excel 파일을 업로드하면 검체를 자동으로 학부별·배양 유형별로 "
        "분류하고 랙 위치를 부여합니다. 바코드 스캐너로 스캔하면 음성(TTS)으로 "
        "위치를 안내하여 분류 오류를 최소화합니다."
    ))

    add_para(doc, "1.2 도입 배경 및 목적", bold=True, size=11)
    add_table(doc,
        ["항목", "내용"],
        [
            ("도입 배경", "수작업 검체 분류로 인한 오류(랙 혼동, 누락) 및 업무 지연 개선"),
            ("핵심 목적", "접수번호 바코드 스캔 → 자동 위치 안내 → 작업 표준화"),
            ("적용 환경", "단일 PC 또는 네트워크 공유(서버 1대 + 클라이언트 N대)"),
            ("주 사용자", "임상검사실 전담 직원 (미생물팀, 혈액·화학·지사 검체 담당)"),
            ("운영 시간", "24시간 운영 (야간 19:00 ~ 익일 09:00 별도 세션)"),
        ],
        col_widths=[4, 12],
    )

    doc.add_paragraph()
    add_para(doc, "1.3 주요 기능 목록", bold=True, size=11)
    add_table(doc,
        ["기능", "설명"],
        [
            ("LIS 파일 업로드", "Excel(xls/xlsx) 접수 목록 → DB 일괄 등록, 중복 방지"),
            ("검체 도착 스캔", "바코드 스캔 → 도착 시각 기록, 학부·분류코드 표시"),
            ("미생물 소분류", "배양 유형별 랙·칸 번호 자동 부여, TTS 음성 안내"),
            ("학부별 소분류", "비미생물 검체를 소분류 기준에 따라 랙 배정"),
            ("지사별 검체 확인", "분류코드 선택 → 워크리스트 생성 → 바코드 스캔 확인"),
            ("Urine 출력", "Urine 해당 행 마킹 Excel 다운로드"),
            ("검체 찾기", "접수번호 검색 → 랙·칸 번호, 도착 여부, 검사명 표시"),
            ("누락검체 조회", "도착하지 않은 검체 목록 표시"),
            ("미접수검체 조회", "LIS에 없는 접수번호 스캔 목록 표시"),
            ("대시보드", "실시간 도착률, 학부별 진행 현황"),
            ("TTS 음성 안내", "Web Speech API, 음성·속도 선택"),
            ("네트워크 공유", "bat 1개로 다른 PC에서 접속"),
        ],
        col_widths=[4.5, 11.5],
    )

    doc.add_page_break()

    # 2. 기술 스택
    add_heading(doc, "2. 기술 스택 및 아키텍처")
    add_table(doc,
        ["구분", "기술", "버전/비고"],
        [
            ("Backend",  "FastAPI",              "0.115 — 비동기 REST API"),
            ("ORM",      "SQLAlchemy",           "2.0 — 선언형 매핑"),
            ("DB",       "SQLite (WAL 모드)",     "파일 기반, 다중 읽기 동시성"),
            ("Frontend", "Jinja2 + Bootstrap 5", "서버 사이드 렌더링"),
            ("음성",     "Web Speech API",        "브라우저 내장 TTS"),
            ("Excel",    "openpyxl / xlrd",       "xlsx/xls 처리"),
            ("런타임",   "Python 3.11 embedded",  "별도 설치 불필요"),
            ("CI/CD",    "GitHub Actions",        "태그 푸시 → 릴리즈 ZIP 자동 생성"),
        ],
        col_widths=[3, 5, 8],
    )
    doc.add_paragraph()
    add_para(doc,
        "아키텍처: FastAPI(uvicorn) 서버 ↔ SQLite DB. 여러 클라이언트 브라우저에서 동시 접속. "
        "3초 주기 폴링으로 실시간 상태 동기화.",
    )

    doc.add_page_break()

    # 3. DB 구조
    add_heading(doc, "3. 데이터베이스 구조 (주요 테이블)")
    add_table(doc,
        ["테이블명", "역할"],
        [
            ("import_batches",                   "업로드 배치 이력"),
            ("orders",                           "검체 접수 목록"),
            ("order_tests",                      "검사 항목 (접수건당 N개)"),
            ("specimen_arrivals",                "검체 도착 기록"),
            ("scan_logs",                        "전체 스캔 이력"),
            ("micro_culture_plans",              "미생물 배양 예정 목록 (업로드 시 생성)"),
            ("micro_culture_assignments",        "미생물 배양 배정 결과"),
            ("department_subcategory_assignments","학부별 소분류 배정 결과"),
            ("branch_rack_sessions",             "지사별 검체 확인 세션 (PC별)"),
            ("branch_rack_items",                "세션 내 검체 (랙번호, 칸번호, 스캔 여부)"),
            ("app_users",                        "사용자 계정"),
            ("app_settings",                     "전역 설정 (출근 기준 시각 등)"),
        ],
        col_widths=[5.5, 10.5],
    )

    doc.add_page_break()

    # 4. 화면별 스펙
    add_heading(doc, "4. 화면별 상세 스펙")

    screens = [
        ("4.1 접수리스트 업로드 (/upload)", [
            ("경로", "/upload"),
            ("목적", "LIS Excel 파일을 DB에 일괄 등록"),
            ("주요 기능", "차수 선택(1차/2차/3차), 중복 접수번호 자동 스킵, 최종 업로드 표시"),
            ("특이사항", "업로드 시 미생물 예정 목록(MicroCulturePlan) 자동 생성\n'파일 삭제 및 초기화' → 지사별 세션 포함 전체 삭제"),
        ]),
        ("4.2 검체 도착·조회 (/scan)", [
            ("경로", "/scan"),
            ("목적", "바코드 스캔으로 검체 도착 기록 및 학부 표시"),
            ("주요 기능", "15자리 → 7자리 자동 변환, 미접수검체 별도 기록"),
        ]),
        ("4.3 미생물 소분류 (/micro)", [
            ("경로", "/micro → 미생물 탭"),
            ("목적", "배양 유형별 랙·칸 자동 배정, TTS 안내"),
            ("특이사항", "업로드 시 접수번호 오름차순 사전 순번 확정\n출근 기준 시각 이후 첫 스캔 시 자동 리셋"),
        ]),
        ("4.4 학부별 소분류 (/micro)", [
            ("경로", "/micro → 비미생물 탭"),
            ("목적", "혈액·화학·요경검 등 비미생물 검체 랙 배정"),
            ("위치 코드", "prefix + sequence_no (예: BCH-001, HEM-003)"),
        ]),
        ("4.5 지사별 검체 확인 (/branch-rack)", [
            ("경로", "/branch-rack"),
            ("목적", "분류코드별 검체를 랙에 배치하고 바코드 스캔으로 확인"),
            ("준비 시작", "동일 코드 재호출 → 신규만 추가 / 코드 변경 → 세션 재생성"),
            ("랙 계산", "rack_no = (sort_no-1)÷50+1  /  rack_position = (sort_no-1)%50+1"),
            ("TTS", "'N번랙 M번' (접수번호 생략)"),
            ("미확인 목록", "접이식 패널 표시, 인쇄 가능 (열: 순번·접수번호·랙·자리·환자명·검체명·병원명)"),
            ("폴링", "3초 주기 자동 갱신"),
        ]),
        ("4.6 Urine 출력 (/urine)", [
            ("경로", "/urine"),
            ("목적", "Urine Random + OC&S 해당 행 마킹 Excel 다운로드"),
        ]),
        ("4.7 검체 찾기 (/find)", [
            ("경로", "/find"),
            ("목적", "접수번호 검색 → 소분류 위치·도착 여부·검사명 한눈에 확인"),
        ]),
        ("4.8 대시보드 (/dashboard)", [
            ("경로", "/dashboard"),
            ("목적", "금일 업무 현황 요약 (도착률, 소분류 진행 현황)"),
        ]),
    ]

    for title_text, rows in screens:
        add_para(doc, title_text, bold=True, size=11)
        add_table(doc, ["항목", "내용"], rows, col_widths=[3.5, 12.5])
        doc.add_paragraph()

    doc.add_page_break()

    # 5. 구현 방향
    add_heading(doc, "5. 구현 방향 및 설계 원칙")
    principles_docx = [
        ("무설치 실행", "Python 3.11 임베디드 런타임 동봉 — 파이썬 설치 없이 bat 더블클릭으로 실행"),
        ("단순 배포", "GitHub Actions로 태그 푸시 시 설치/업데이트 ZIP 2종 자동 생성"),
        ("데이터 영속성", "SQLite WAL 모드로 다중 브라우저 동시 읽기 지원, DB 파일 1개가 전체 백업"),
        ("누적 업로드", "1차→2차→3차 순차 업로드 시 신규 접수번호만 추가, 기존 스캔 기록 보존"),
        ("랙 위치 사전 부여", "업로드 즉시 접수번호 오름차순으로 MicroCulturePlan 순번 확정"),
        ("폴링 기반 실시간성", "3초 주기 REST API 폴링으로 다중 PC 간 상태 동기화"),
        ("TTS 표준화", "Web Speech API, 화면 간 음성·속도 설정 공유 (localStorage)"),
        ("서비스 레이어 분리", "라우터(HTTP 처리) ↔ 서비스(비즈니스 로직) 분리"),
        ("점진적 스키마", "서버 시작 시 ensure_schema_columns()로 누락 컬럼 ALTER TABLE 자동 추가"),
        ("보안", "bcrypt 비밀번호 해싱, 관리자 승인 후 로그인 허용"),
    ]
    add_table(doc,
        ["원칙", "내용"],
        principles_docx,
        col_widths=[4, 12],
    )

    doc.add_page_break()

    # 6. API 목록
    add_heading(doc, "6. API 엔드포인트 목록")
    add_table(doc,
        ["메서드", "경로", "설명"],
        [
            ("POST", "/api/import/orders",            "Excel 업로드 → 접수 목록 등록"),
            ("POST", "/api/import/reset-all",         "전체 데이터 초기화"),
            ("POST", "/api/import/urine-mark",        "Urine 마킹 Excel 다운로드"),
            ("GET",  "/api/import/worklist",          "워크리스트 Excel 다운로드"),
            ("POST", "/api/scans/arrive",             "검체 도착 스캔"),
            ("GET",  "/api/scans/summary",            "도착 현황 요약"),
            ("POST", "/api/micro/scan",               "미생물 소분류 스캔"),
            ("GET",  "/api/micro/session",            "현재 미생물 세션 조회"),
            ("POST", "/api/micro/reset",              "미생물 세션 리셋"),
            ("POST", "/api/micro/subdivision/scan",   "학부별 소분류 스캔"),
            ("GET",  "/api/micro/subdivision/session","학부별 소분류 세션 조회"),
            ("GET",  "/api/reports/missing",          "누락검체 목록"),
            ("GET",  "/api/reports/unregistered",     "미접수검체 목록"),
            ("GET",  "/api/reports/dashboard",        "대시보드 통계"),
            ("POST", "/api/branch-rack/prepare",      "지사별 검체 확인 세션 준비"),
            ("GET",  "/api/branch-rack/session",      "지사별 세션 조회"),
            ("POST", "/api/branch-rack/scan",         "지사별 바코드 스캔"),
            ("GET",  "/api/branch-rack/codes",        "사용 가능한 분류코드 목록"),
            ("POST", "/api/auth/register",            "회원가입"),
            ("POST", "/api/auth/login",               "로그인"),
            ("GET",  "/api/auth/users",               "사용자 목록 (관리자)"),
            ("POST", "/api/auth/approve/{id}",        "사용자 승인 (관리자)"),
        ],
        col_widths=[2, 7, 7],
    )

    doc.add_page_break()

    # 7. 사용설명서
    add_heading(doc, "7. 화면별 사용설명서")

    add_para(doc, "7.1 시작 및 로그인", bold=True, size=11)
    add_table(doc,
        ["단계", "내용"],
        [
            ("시작",       "시작.bat (또는 검체관리프로그램.exe) 더블클릭 → 브라우저 자동 오픈"),
            ("주소",       "http://127.0.0.1:8000 (로컬) 또는 http://서버IP:8000 (네트워크)"),
            ("로그인",     "아이디 / 비밀번호 입력 (초기 계정: admin1234 / admin1234)"),
            ("신규 계정",  "회원가입 → 관리자(/admin/users)에서 승인 후 사용 가능"),
        ],
        col_widths=[3.5, 12.5],
    )
    doc.add_paragraph()

    add_para(doc, "7.2 접수리스트 업로드", bold=True, size=11)
    steps_ul = [
        "① 접수리스트 업로드 화면으로 이동",
        "② 파일 선택 → LIS Excel 파일(.xls/.xlsx) 선택",
        "③ 차수 선택: 당일 처음이면 1차, 추가 접수가 있으면 2차·3차",
        "④ 오늘 마지막 업로드라면 '최종 업로드' 체크",
        "⑤ 평일/토요일 선택 후 업로드 버튼 클릭",
        "⑥ 성공 시 '○○건 등록 완료' 토스트 표시",
        "※ 동일 접수번호 중복 업로드 시 해당 건은 자동 스킵",
    ]
    for s in steps_ul:
        add_para(doc, s, indent=0.5, color=(0x92, 0x40, 0x0e) if s.startswith("※") else None)
    doc.add_paragraph()

    add_para(doc, "7.3 미생물 소분류 스캔", bold=True, size=11)
    steps_micro = [
        "① /micro 화면 이동 → 미생물 탭 선택",
        "② 배양 유형 버튼 선택 (복수 가능 — 예: 일반세균 + 혈액배양)",
        "③ TTS 토글 ON → 음성 안내 활성화",
        "④ 바코드 스캔 → '2번랙 15번칸' 음성 + 화면 표시",
        "⑤ 랙 그리드에서 주황 점멸 위치를 확인 후 검체 꽂기",
        "⑥ 랙 가득 차면 다음 랙으로 자동 이동",
        "※ 출근 기준 시각(기본 19:00) 이후 첫 스캔 시 랙 자동 리셋",
    ]
    for s in steps_micro:
        add_para(doc, s, indent=0.5, color=(0x92, 0x40, 0x0e) if s.startswith("※") else None)
    doc.add_paragraph()

    add_para(doc, "7.4 지사별 검체 확인", bold=True, size=11)
    steps_br = [
        "① /branch-rack 화면 이동",
        "② PC명 입력 (예: 지사1_PC, 야간검사실_PC)",
        "③ 분류코드 버튼 선택 (접수번호 앞 2자리 — 복수 선택 가능)",
        "④ '준비 시작' 클릭 → 해당 코드 검체 목록 생성 + 랙 배정",
        "⑤ 바코드 스캔 → '1번랙 5번' 음성 + 랙 그리드 강조",
        "⑥ '미확인 N건' 클릭 → 미스캔 목록 펼침 (인쇄 가능)",
        "※ 동일 코드로 재클릭 시 신규 검체만 추가, 기존 스캔 상태 유지",
        "※ 코드 변경 시 세션 초기화 후 새로 생성",
    ]
    for s in steps_br:
        add_para(doc, s, indent=0.5, color=(0x92, 0x40, 0x0e) if s.startswith("※") else None)
    doc.add_paragraph()

    add_para(doc, "7.5 네트워크 공유", bold=True, size=11)
    add_table(doc,
        ["항목", "내용"],
        [
            ("서버 측",   "네트워크공유.bat 실행 → 0.0.0.0:8000 바인딩"),
            ("클라이언트", "브라우저에서 http://서버IP:8000 접속"),
            ("방화벽",    "서버 PC에서 TCP 8000 인바운드 허용 필요"),
            ("동시 접속", "복수 PC 동시 접속, 3초 폴링으로 상태 공유"),
        ],
        col_widths=[3.5, 12.5],
    )

    doc.add_page_break()

    # 8. 배포
    add_heading(doc, "8. 배포 및 설치 안내")
    add_para(doc, "8.1 배포 파일 종류", bold=True, size=11)
    add_table(doc,
        ["파일명", "대상", "내용"],
        [
            ("검체관리프로그램_설치_vX.X.X.zip", "처음 설치", "전체 파일 (python 포함, 약 65 MB)"),
            ("검체관리프로그램_업데이트_vX.X.X.zip", "업데이트", "app 폴더만 — 기존 경로에 덮어쓰기"),
        ],
        col_widths=[6, 3.5, 6.5],
    )
    doc.add_paragraph()

    add_para(doc, "8.2 신규 설치", bold=True, size=11)
    for s in [
        "① 설치 ZIP 다운로드 및 압축 해제",
        "② 시작.bat 더블클릭 → 브라우저 자동 오픈",
        "③ admin1234 / admin1234 로 첫 로그인",
    ]:
        add_para(doc, s, indent=0.5)
    doc.add_paragraph()

    add_para(doc, "8.3 업데이트", bold=True, size=11)
    for s in [
        "① 업데이트 ZIP 다운로드",
        "② 기존 서버 종료",
        "③ 업데이트 ZIP의 app 폴더를 기존 경로에 덮어쓰기",
        "④ 시작.bat 재실행 (DB 파일은 유지)",
    ]:
        add_para(doc, s, indent=0.5)
    doc.add_paragraph()

    add_para(doc, "8.4 데이터 백업", bold=True, size=11)
    add_para(doc,
        "specimen_routing.db 파일 하나를 복사하면 완료. "
        "서버 실행 중에도 복사 가능 (WAL 모드).")

    doc.add_page_break()

    # 9. GitHub
    add_heading(doc, "9. GitHub 관리 안내")
    add_para(doc,
        "소스 코드 변경 후 아래 순서로 커밋·태그 푸시하면 GitHub Actions가 "
        "배포 ZIP 2종을 자동 생성합니다.")
    add_table(doc,
        ["단계", "명령어"],
        [
            ("변경 확인", "git status"),
            ("스테이징", "git add -A"),
            ("커밋", 'git commit -m "feat: 변경 내용"'),
            ("푸시", "git push origin main"),
            ("버전 태그", "git tag v1.5.0"),
            ("태그 푸시 → 릴리즈 자동 생성", "git push origin v1.5.0"),
        ],
        col_widths=[5, 11],
    )
    doc.add_paragraph()
    add_para(doc,
        "로컬 배포 ZIP이 필요하면 빌드_배포파일.ps1 (또는 빌드_배포파일.bat)을 "
        "실행하세요. 결과는 배포\\ 폴더에 날짜 포함 파일명으로 저장됩니다.",
        color=(0x6b, 0x72, 0x80),
    )

    doc.add_paragraph()
    add_para(doc, f"본 문서는 {TODAY} 기준으로 생성되었습니다.",
             color=(0x6b, 0x72, 0x80), size=9)

    doc.save(str(DOCX_PATH))
    print(f"DOCX 생성 완료: {DOCX_PATH}")


if __name__ == "__main__":
    print("PDF 생성 중...")
    build_pdf()
    print("DOCX 생성 중...")
    build_docx()
    print("\n모든 문서 생성 완료!")
    print(f"  PDF : {PDF_PATH}")
    print(f"  DOCX: {DOCX_PATH}")
